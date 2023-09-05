from __future__ import print_function
from flask import Flask, render_template, request, redirect, make_response, session
from docx import Document
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import zipfile
from datetime import datetime
import locale
from urllib.parse import quote
from flask import url_for
import boto3
from botocore.exceptions import NoCredentialsError
import os.path
from utils.date_utils import format_data_extenso
from utils.upload_s3 import upload_to_s3

app = Flask(__name__)

# Configurações de autenticação
SCOPES = ['https://www.googleapis.com/auth/spreadsheets']


@app.route('/')
def index():
    return render_template('index.html')

AWS_ACCESS_KEY = os.environ.get('AWS_ACCESS_KEY')
AWS_SECRET_KEY = os.environ.get('AWS_SECRET_KEY')


@app.route('/generate_docx', methods=['POST'])
def gerar_docx():
    try:
        nome = request.form['nome'].upper()
        nacionalidade = request.form['nacionalidade'].upper()
        estadoCivil = request.form['estadoCivil'].upper()
        profissao = request.form['profissao'].upper()
        fone = request.form['fone'].upper()
        fone_recado = request.form['fone_recado'].upper()
        cpf = request.form['cpf'].upper()
        rg = request.form['rg'].upper()
        data_nascimento = request.form['data_nascimento']
        endereco = request.form['endereco'].upper()
        bairro = request.form['bairro'].upper()
        cep = request.form['cep'].upper()
        cidade = request.form['cidade'].upper()
        estado = request.form['estado'].upper()
        cep = request.form['cep'].upper()
        data_str = request.form['data']

        # Chamando a função format_data_extenso para obter a data por extenso
        data_extenso = format_data_extenso(data_str)

        # Converta a data para o formato de data do Python
        data_nascimento_str = datetime.strptime(data_nascimento, "%Y/%m/%d")
        # Formate a data como "dia/mês/ano" (ou o formato desejado)
        data_formatada = data_nascimento_str.strftime("%d/%m/%Y")
        
        
        ### DOC1 = CONTRATO HONORARIOS ###
        doc1 = Document('./modelos/contratoHonorarios.docx') # Substitua 'modelo.docx' pelo caminho do seu modelo DOCX
        for table in doc1.tables: #percorrendo todas as tabelas
            for row in table.rows: #percorrendo todas as linhas
                for cell in row.cells: #percorrendo todas as cedulas 
                    cell_text = cell.text 
                    if '{{nome}}' in cell_text: #se encontrar {{NOME}}
                        cell.text = cell_text.replace('{{nome}}', nome) #vai inserir o valor do usuario
                    if '{{nacionalidade}}' in cell_text:
                        cell.text = cell_text.replace('{{nacionalidade}}', nacionalidade)
                    if '{{estadoCivil}}' in cell_text:
                        cell.text = cell_text.replace('{{estadoCivil}}', estadoCivil)
                    if '{{profissao}}' in cell_text:
                        cell.text = cell_text.replace('{{profissao}}', profissao)
                    if '{{fone}}' in cell_text:
                        if fone_recado:
                            cell.text = cell_text.replace('{{fone}}', fone + f' ou {fone_recado}')
                        else:
                            cell.text = cell_text.replace('{{fone}}', fone)
                    if '{{cpf}}' in cell_text:
                        cell.text = cell_text.replace('{{cpf}}', cpf)
                    if '{{rg}}' in cell_text:
                        cell.text = cell_text.replace('{{rg}}', rg)
                    if '{{endereco}}' in cell_text:
                        cell.text = cell_text.replace('{{endereco}}', endereco)
                    if '{{bairro}}' in cell_text:
                        cell.text = cell_text.replace('{{bairro}}', bairro)
                    if '{{cep}}' in cell_text:
                        cell.text = cell_text.replace('{{cep}}', cep)
                    if '{{cidade}}' in cell_text:
                        cell.text = cell_text.replace('{{cidade}}', cidade)
                    if '{{estado}}' in cell_text:
                        cell.text = cell_text.replace('{{estado}}', estado)
        
        for paragraph in doc1.paragraphs: #percorre os paragratos
            paragraph_text = paragraph.text
            if '{{data}}' in paragraph_text: #encontrando {{data}} substitui
                paragraph.text = paragraph_text.replace('{{data}}', data_extenso)

        doc1_path = os.path.join('modelos', f'CONTRATO HONORÁRIO - {nome}.docx')
        doc1.save(doc1_path)

        ### DOC2 = JUSTIÇA GRATUIDA ###
        doc2 = Document('./modelos/justicagratuita.docx')
        for table in doc2.tables: #percorrendo todas as tabelas
            for row in table.rows: #percorrendo todas as linhas
                for cell in row.cells: #percorrendo todas as cedulas 
                    cell_text = cell.text 
                    if '{{nome}}' in cell_text: #se encontrar {{NOME}}
                        cell.text = cell_text.replace('{{nome}}', nome) #vai inserir o valor do usuario
                    if '{{nacionalidade}}' in cell_text:
                        cell.text = cell_text.replace('{{nacionalidade}}', nacionalidade)
                    if '{{estadoCivil}}' in cell_text:
                        cell.text = cell_text.replace('{{estadoCivil}}', estadoCivil)
                    if '{{profissao}}' in cell_text:
                        cell.text = cell_text.replace('{{profissao}}', profissao)
                    if '{{fone}}' in cell_text:
                        if fone_recado:
                            cell.text = cell_text.replace('{{fone}}', fone + f' ou {fone_recado}')
                        else:
                            cell.text = cell_text.replace('{{fone}}', fone)
                    if '{{cpf}}' in cell_text:
                        cell.text = cell_text.replace('{{cpf}}', cpf)
                    if '{{rg}}' in cell_text:
                        cell.text = cell_text.replace('{{rg}}', rg)
                    if '{{endereco}}' in cell_text:
                        cell.text = cell_text.replace('{{endereco}}', endereco)
                    if '{{bairro}}' in cell_text:
                        cell.text = cell_text.replace('{{bairro}}', bairro)
                    if '{{cep}}' in cell_text:
                        cell.text = cell_text.replace('{{cep}}', cep)
                    if '{{cidade}}' in cell_text:
                        cell.text = cell_text.replace('{{cidade}}', cidade)
                    if '{{estado}}' in cell_text:
                        cell.text = cell_text.replace('{{estado}}', estado)
        
        for paragraph in doc2.paragraphs: #percorre os paragratos
            paragraph_text = paragraph.text
            if '{{data}}' in paragraph_text: #encontrando {{data}} substitui
                paragraph.text = paragraph_text.replace('{{data}}', data_extenso)

        doc2_path = os.path.join('modelos', f'JUSTIÇA GRATUITA - {nome}.docx')
        doc2.save(doc2_path)

        ### DOC3 = PROCURAÇÃO ###
        doc3 = Document('./modelos/procuracao.docx')
        for table in doc3.tables: #percorrendo todas as tabelas
            for row in table.rows: #percorrendo todas as linhas
                for cell in row.cells: #percorrendo todas as cedulas 
                    cell_text = cell.text 
                    if '{{nome}}' in cell_text: #se encontrar {{NOME}}
                        cell.text = cell_text.replace('{{nome}}', nome) #vai inserir o valor do usuario
                    if '{{nacionalidade}}' in cell_text:
                        cell.text = cell_text.replace('{{nacionalidade}}', nacionalidade)
                    if '{{estadoCivil}}' in cell_text:
                        cell.text = cell_text.replace('{{estadoCivil}}', estadoCivil)
                    if '{{profissao}}' in cell_text:
                        cell.text = cell_text.replace('{{profissao}}', profissao)
                    if '{{fone}}' in cell_text:
                        if fone_recado:
                            cell.text = cell_text.replace('{{fone}}', fone + f' ou {fone_recado}')
                        else:
                            cell.text = cell_text.replace('{{fone}}', fone)
                    if '{{cpf}}' in cell_text:
                        cell.text = cell_text.replace('{{cpf}}', cpf)
                    if '{{rg}}' in cell_text:
                        cell.text = cell_text.replace('{{rg}}', rg)
                    if '{{endereco}}' in cell_text:
                        cell.text = cell_text.replace('{{endereco}}', endereco)
                    if '{{bairro}}' in cell_text:
                        cell.text = cell_text.replace('{{bairro}}', bairro)
                    if '{{cep}}' in cell_text:
                        cell.text = cell_text.replace('{{cep}}', cep)
                    if '{{cidade}}' in cell_text:
                        cell.text = cell_text.replace('{{cidade}}', cidade)
                    if '{{estado}}' in cell_text:
                        cell.text = cell_text.replace('{{estado}}', estado)
        
        for paragraph in doc3.paragraphs: #percorre os paragratos
            paragraph_text = paragraph.text
            if '{{data}}' in paragraph_text: #encontrando {{data}} substitui
                paragraph.text = paragraph_text.replace('{{data}}', data_extenso)

        doc3_path = os.path.join('modelos', f'PROCURAÇÃO - {nome}.docx')
        doc3.save(doc3_path)
       
        ### DOC4 = Capa Processo ###
        doc4 = Document('./modelos/capaProcesso.docx')
        for table in doc4.tables: #percorrendo todas as tabelas
            for row in table.rows: #percorrendo todas as linhas
                for cell in row.cells: #percorrendo todas as cedulas 
                    cell_text = cell.text 
                    if '{{nome}}' in cell_text: #se encontrar {{NOME}}
                        cell.text = cell_text.replace('{{nome}}', nome) #vai inserir o valor do usuario
                    if '{{nacionalidade}}' in cell_text:
                        cell.text = cell_text.replace('{{nacionalidade}}', nacionalidade)
                    if '{{estadoCivil}}' in cell_text:
                        cell.text = cell_text.replace('{{estadoCivil}}', estadoCivil)
                    if '{{profissao}}' in cell_text:
                        cell.text = cell_text.replace('{{profissao}}', profissao)
                    if '{{fone}}' in cell_text:
                        if fone_recado:
                            cell.text = cell_text.replace('{{fone}}', fone + f' ou {fone_recado}')
                        else:
                            cell.text = cell_text.replace('{{fone}}', fone)
                    if '{{cpf}}' in cell_text:
                        cell.text = cell_text.replace('{{cpf}}', cpf)
                    if '{{rg}}' in cell_text:
                        cell.text = cell_text.replace('{{rg}}', rg)
                    if '{{endereco}}' in cell_text:
                        cell.text = cell_text.replace('{{endereco}}', endereco)
                    if '{{bairro}}' in cell_text:
                        cell.text = cell_text.replace('{{bairro}}', bairro)
                    if '{{cep}}' in cell_text:
                        cell.text = cell_text.replace('{{cep}}', cep)
                    if '{{cidade}}' in cell_text:
                        cell.text = cell_text.replace('{{cidade}}', cidade)
                    if '{{estado}}' in cell_text:
                        cell.text = cell_text.replace('{{estado}}', estado)

        doc4_path = os.path.join('modelos', f'CAPA DO PROCESSO - {nome}.docx')
        doc4.save(doc4_path)


        ### DOC5 = Minuta Auxilio Acidente Federal ###
        doc5 = Document('./modelos/minutaAuxilioAcidenteFederal.docx')
           
        for paragraph in doc5.paragraphs:
            paragraph_text = paragraph.text
            if '{{nome}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{nome}}', nome)
            if '{{nacionalidade}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{nacionalidade}}', nacionalidade) 
            if '{{estadoCivil}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{estadoCivil}}', estadoCivil)
            if '{{profissao}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{profissao}}', profissao)
            if '{{cpf}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cpf}}', cpf)
            if '{{rg}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{rg}}', rg)
            if '{{endereco}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{endereco}}', endereco)
            if '{{bairro}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{bairro}}', bairro)
            if '{{cidade}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cidade}}', cidade)
            if '{{estado}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{estado}}', estado)
            if '{{cep}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cep}}', cep)
            if '{{data}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{data}}', data_extenso)
   
            # Atribuir o texto modificado de volta ao parágrafo
            paragraph.text = paragraph_text

        doc5_path = os.path.join('modelos', f'MINUTA AUXILIO ACIDENTE FEDERAL - {nome}.docx')
        doc5.save(doc5_path)


        ### DOC6 = Requerimento Adm Auxilio Acidente ###
        doc6 = Document('./modelos/requerimentoAdmAuxilioAcidente.docx')

        for paragraph in doc6.paragraphs:
            paragraph_text = paragraph.text
            if '{{nome}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{nome}}', nome)
            if '{{nacionalidade}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{nacionalidade}}', nacionalidade)
            if '{{estadoCivil}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{estadoCivil}}', estadoCivil)
            if '{{profissao}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{profissao}}', profissao)
            if '{{cpf}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cpf}}', cpf)
            if '{{rg}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{rg}}', rg)
            if '{{endereco}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{endereco}}', endereco)
            if '{{bairro}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{bairro}}', bairro)
            if '{{cidade}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cidade}}', cidade)
            if '{{estado}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{estado}}', estado)
            if '{{cep}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cep}}', cep)
            if '{{data}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{data}}', data_extenso)

            # Atribuir o texto modificado de volta ao parágrafo
            paragraph.text = paragraph_text

        doc6_path = os.path.join('modelos', f'REQUERIMENTO ADMINISTRATIVO AUXILIO ACIDENTE - {nome}.docx')
        doc6.save(doc6_path)

        ### DOC7 = Declaração de Residencia ###
        doc7 = Document('./modelos/declaracaoDeResidencia.docx')
        for paragraph in doc7.paragraphs:
            paragraph_text = paragraph.text
            if '{{nome}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{nome}}', nome) 
            if '{{cpf}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cpf}}', cpf)
            if '{{rg}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{rg}}', rg)
            if '{{fone}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{fone}}', fone)
            if '{{endereco}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{endereco}}', endereco)
            if '{{bairro}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{bairro}}', bairro)
            if '{{cidade}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cidade}}', cidade)
            if '{{estado}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{estado}}', estado)
            if '{{data}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{data}}', data_extenso)

            # Atribuir o texto modificado de volta ao parágrafo
            paragraph.text = paragraph_text

        doc7_path = os.path.join('modelos', f'DECLARAÇÃO DE RESIDENCIA - {nome}.docx')
        doc7.save(doc7_path)
        
        ### DOC8 = Termo de Renúncia ###
        doc8 = Document('./modelos/termoDeRenuncia.docx')
        for table in doc8.tables: #percorrendo todas as tabelas
            for row in table.rows: #percorrendo todas as linhas
                for cell in row.cells: #percorrendo todas as cedulas 
                    cell_text = cell.text 
                    if '{{nome}}' in cell_text: #se encontrar {{NOME}}
                        cell.text = cell_text.replace('{{nome}}', nome) #vai inserir o valor do usuario
                    if '{{nacionalidade}}' in cell_text:
                        cell.text = cell_text.replace('{{nacionalidade}}', nacionalidade)
                    if '{{estadoCivil}}' in cell_text:
                        cell.text = cell_text.replace('{{estadoCivil}}', estadoCivil)
                    if '{{profissao}}' in cell_text:
                        cell.text = cell_text.replace('{{profissao}}', profissao)
                    if '{{fone}}' in cell_text:
                        if fone_recado:
                            cell.text = cell_text.replace('{{fone}}', fone + f' ou {fone_recado}')
                        else:
                            cell.text = cell_text.replace('{{fone}}', fone)
                    if '{{cpf}}' in cell_text:
                        cell.text = cell_text.replace('{{cpf}}', cpf)
                    if '{{rg}}' in cell_text:
                        cell.text = cell_text.replace('{{rg}}', rg)
                    if '{{endereco}}' in cell_text:
                        cell.text = cell_text.replace('{{endereco}}', endereco)
                    if '{{bairro}}' in cell_text:
                        cell.text = cell_text.replace('{{bairro}}', bairro)
                    if '{{cep}}' in cell_text:
                        cell.text = cell_text.replace('{{cep}}', cep)
                    if '{{cidade}}' in cell_text:
                        cell.text = cell_text.replace('{{cidade}}', cidade)
                    if '{{estado}}' in cell_text:
                        cell.text = cell_text.replace('{{estado}}', estado)

        for paragraph in doc8.paragraphs: #percorre os paragratos
            paragraph_text = paragraph.text
            if '{{cidade}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{cidade}}', cidade)
            if '{{data}}' in paragraph_text: 
                paragraph_text = paragraph_text.replace('{{data}}', data_extenso)
            if '{{nome}}' in paragraph_text:
                paragraph_text = paragraph_text.replace('{{nome}}', nome) 
            paragraph.text = paragraph_text
        
        doc8_path = os.path.join('modelos', f'TERMO DE RENÚNCIA - {nome}.docx')
        doc8.save(doc8_path)
        
        # Fazer upload dos documentos para o S3
        if upload_to_s3(doc1_path, 'cadastroadv', f'datas/CONTRATO HONORÁRIO - {nome}.docx') and \
           upload_to_s3(doc2_path, 'cadastroadv', f'datas/JUSTIÇA GRATUITA - {nome}.docx') and \
           upload_to_s3(doc3_path, 'cadastroadv', f'datas/PROCURAÇÃO - {nome}.docx') and \
           upload_to_s3(doc4_path, 'cadastroadv', f'datas/CAPA DO PROCESSO - {nome}.docx') and \
           upload_to_s3(doc5_path, 'cadastroadv', f'datas/MINUTA AUXILIO ACIDENTE FEDERAL - {nome}.docx') and \
           upload_to_s3(doc6_path, 'cadastroadv', f'datas/REQUERIMENTO ADMINISTRATIVO AUXILIO ACIDENTE - {nome}.docx') and \
           upload_to_s3(doc7_path, 'cadastroadv', f'datas/DECLARAÇÃO DE RESIDENCIA - {nome}.docx') and \
           upload_to_s3(doc8_path, 'cadastroadv', f'datas/TERMO DE RENÚNCIA - {nome}.docx'):
           return redirect(url_for('download_files', 
                                   nome=nome, 
                                   estadoCivil=estadoCivil,
                                   profissao=profissao, 
                                   fone=fone,
                                   fone_recado=fone_recado,
                                   cpf=cpf,
                                   rg=rg,
                                   data_nascimento=data_formatada,
                                   endereco=endereco,
                                   bairro=bairro, 
                                   cidade=cidade,
                                   estado=estado,
                                   cep=cep
                                   ))
           #return "Documentos gerados e enviados com sucesso!"
        else:
            return f"Erro ao gerar e/ou enviar os documentos."
        
        return "ok"
    except KeyError as e:
        return f"Erro: O campo '{e.args[0]}' não foi encontrado nos dados enviados."
    except Exception as e:
        return f"Erro inesperado: {str(e)}"



@app.route('/downloads/<nome>')
def download_files(nome):
    s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

    filenames = [f'CONTRATO HONORÁRIO - {nome}.docx', 
                    f'JUSTIÇA GRATUITA - {nome}.docx', 
                    f'PROCURAÇÃO - {nome}.docx', 
                    f'CAPA DO PROCESSO - {nome}.docx', 
                    f'MINUTA AUXILIO ACIDENTE FEDERAL - {nome}.docx',
                    f'REQUERIMENTO ADMINISTRATIVO AUXILIO ACIDENTE - {nome}.docx',
                    f'DECLARAÇÃO DE RESIDENCIA - {nome}.docx',
                    f'TERMO DE RENÚNCIA - {nome}.docx'
                ]
    
    download_links = []

    estadoCivil = request.args.get('estadoCivil')
    profissao = request.args.get('profissao')
    fone = request.args.get('fone')
    fone_recado = request.args.get('fone_recado')
    cpf = request.args.get('cpf')
    rg = request.args.get('rg')
    data_nascimento = request.args.get('data_formatada')
    endereco = request.args.get('endereco')
    bairro = request.args.get('bairro')
    cidade = request.args.get('cidade')
    estado = request.args.get('estado')
    cep = request.args.get('cep')


    try:
        for filename in filenames:
            url = s3.generate_presigned_url('get_object',
                                           Params={'Bucket': 'cadastroadv', 'Key': f'datas/{filename}'},
                                           ExpiresIn=3600)
            download_links.append({'filename': filename, 'download_link': url})
            
        return render_template('download.html', download_links=download_links, 
                               nome=nome,
                               estadoCivil=estadoCivil,
                               profissao=profissao, 
                               fone=fone,
                               fone_recado=fone_recado,
                               cpf=cpf,
                               rg=rg,
                               data_nascimento=data_nascimento,
                               endereco=endereco,
                               bairro=bairro,
                               cidade=cidade,
                               estado=estado,
                               cep=cep
                               )
    except NoCredentialsError:
        return "Credenciais do AWS não foram configuradas."



if __name__ == '__main__':
    #app.run(debug=True)
    app.run(host='0.0.0.0', port=os.environ.get('PORT', 5000))